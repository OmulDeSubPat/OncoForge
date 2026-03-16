from __future__ import annotations

import json
from collections import Counter
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH

from src.config import PROJECT_ROOT
from src.knowledge import BUZZWORD_ENTRIES, BuzzwordEntry


REPORTS_DIR = PROJECT_ROOT / "reports"
DOCX_PATH = REPORTS_DIR / "OncoForge_Buzzword_Glossary.docx"
HISTORY_DIR = REPORTS_DIR / "buzzword_glossary_history"
HISTORY_INDEX = HISTORY_DIR / "build_history.json"
CONTEXT_MEMORY = HISTORY_DIR / "glossary_context_memory.md"

CATEGORY_ORDER = [
    "Project Architecture",
    "Machine Learning",
    "Reinforcement Learning",
    "Cheminformatics",
    "Medicinal Chemistry",
    "Structural Biology",
    "Cancer Biology",
]


def _load_json(path: Path) -> dict | list | None:
    if not path.exists():
        return None
    return json.loads(path.read_text(encoding="utf-8"))


def _fmt(value) -> str:
    try:
        return f"{float(value):.3f}"
    except (TypeError, ValueError):
        return "n/a"


def _term_id(term: str) -> str:
    return (
        term.lower()
        .replace("/", "_")
        .replace("-", "_")
        .replace(" ", "_")
        .replace("(", "")
        .replace(")", "")
        .replace(".", "")
    )


def _load_history() -> list[dict]:
    history = _load_json(HISTORY_INDEX)
    return history if isinstance(history, list) else []


def _save_history(history: list[dict]) -> None:
    HISTORY_DIR.mkdir(parents=True, exist_ok=True)
    HISTORY_INDEX.write_text(json.dumps(history, indent=2), encoding="utf-8")


def _project_snapshot() -> dict:
    notebook_metrics = _load_json(REPORTS_DIR / "technical_notebook" / "technical_notebook_metrics.json") or {}
    model_summary = _load_json(REPORTS_DIR / "model_performance_summary.json") or {}
    rl_summary = _load_json(REPORTS_DIR / "rl_verifiable" / "rl_training_summary.json") or {}
    rl_top = rl_summary.get("top_candidate", [])
    return {
        "dataset_name": model_summary.get("dataset_name"),
        "scaffold_rmse": model_summary.get("scaffold_split", {}).get("rmse"),
        "mean_feasibility_score": notebook_metrics.get("mean_feasibility_score"),
        "cross_database_mean_consensus": notebook_metrics.get("cross_database_mean_consensus"),
        "cross_database_strong_rate": notebook_metrics.get("cross_database_strong_rate"),
        "mean_external_evidence_support": notebook_metrics.get("external_evidence_mean_support"),
        "rl_mean_external_evidence_support": notebook_metrics.get("rl_mean_external_evidence_support"),
        "pubchem_mean_enriched_evidence_score": notebook_metrics.get("pubchem_mean_enriched_evidence_score"),
        "best_vina_affinity_kcal": notebook_metrics.get("best_vina_affinity_kcal"),
        "mean_interaction_support": notebook_metrics.get("mean_interaction_support"),
        "audit_pass_rate": notebook_metrics.get("audit_pass_rate"),
        "source_holdout_mean_rmse": notebook_metrics.get("source_holdout_mean_rmse"),
        "rediscovery_protected_top10_recall": notebook_metrics.get("rediscovery_protected_top10_recall"),
        "rl_best_episode_return": rl_summary.get("best_episode_return"),
        "rl_top_candidate": rl_top[0] if isinstance(rl_top, list) and rl_top else None,
    }


def _prepare_snapshot(history: list[dict]) -> dict:
    category_counts = Counter(entry.category for entry in BUZZWORD_ENTRIES)
    current_terms = {_term_id(entry.term) for entry in BUZZWORD_ENTRIES}
    previous_terms = set(history[-1].get("term_ids", [])) if history else set()
    return {
        "run_label": datetime.now().strftime("%Y%m%d_%H%M%S"),
        "created_at": datetime.now().isoformat(timespec="seconds"),
        "term_count": len(BUZZWORD_ENTRIES),
        "used_in_project_count": sum(1 for entry in BUZZWORD_ENTRIES if entry.used_in_project),
        "category_counts": dict(sorted(category_counts.items())),
        "term_ids": sorted(current_terms),
        "new_terms": sorted(current_terms - previous_terms),
        "project_snapshot": _project_snapshot(),
    }


def _append_context_memory(snapshot: dict) -> None:
    HISTORY_DIR.mkdir(parents=True, exist_ok=True)
    project = snapshot.get("project_snapshot", {})
    lines = [
        f"## {snapshot.get('run_label', 'unknown_run')}",
        f"- Created at: {snapshot.get('created_at', 'n/a')}",
        f"- Glossary terms: {snapshot.get('term_count', 0)} total, {snapshot.get('used_in_project_count', 0)} used directly in OncoForge.",
        f"- Scaffold RMSE snapshot: {_fmt(project.get('scaffold_rmse'))}",
        f"- Mean feasibility snapshot: {_fmt(project.get('mean_feasibility_score'))}",
        f"- Cross-database mean consensus snapshot: {_fmt(project.get('cross_database_mean_consensus'))}",
        f"- Cross-database strong rate snapshot: {_fmt(project.get('cross_database_strong_rate'))}",
        f"- External evidence support snapshot: {_fmt(project.get('mean_external_evidence_support'))}",
        f"- Source holdout mean RMSE snapshot: {_fmt(project.get('source_holdout_mean_rmse'))}",
        f"- Rediscovery protected top10 snapshot: {_fmt(project.get('rediscovery_protected_top10_recall'))}",
        f"- RL external evidence snapshot: {_fmt(project.get('rl_mean_external_evidence_support'))}",
        f"- PubChem enriched evidence snapshot: {_fmt(project.get('pubchem_mean_enriched_evidence_score'))}",
        f"- Best Vina snapshot: {_fmt(project.get('best_vina_affinity_kcal'))}",
        f"- RL best episode return snapshot: {_fmt(project.get('rl_best_episode_return'))}",
    ]
    new_terms = snapshot.get("new_terms", [])
    if new_terms:
        lines.append(f"- New terms since previous build: {', '.join(new_terms[:15])}")
    lines.append("")
    with CONTEXT_MEMORY.open("a", encoding="utf-8") as handle:
        handle.write("\n".join(lines) + "\n")


def _add_title(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = document.styles["Title"].font.size


def _add_label(document: Document, label: str, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.add_run(f"{label}: ").bold = True
    paragraph.add_run(text)


def _add_history_table(document: Document, history: list[dict]) -> None:
    if not history:
        return
    document.add_heading("Update History", level=1)
    table = document.add_table(rows=1, cols=6)
    header = table.rows[0].cells
    header[0].text = "Build"
    header[1].text = "Terms"
    header[2].text = "Used Now"
    header[3].text = "Scaffold RMSE"
    header[4].text = "Mean Feasibility"
    header[5].text = "Best Vina"
    for snapshot in history[-8:]:
        row = table.add_row().cells
        project = snapshot.get("project_snapshot", {})
        row[0].text = str(snapshot.get("run_label", "n/a"))
        row[1].text = str(snapshot.get("term_count", "n/a"))
        row[2].text = str(snapshot.get("used_in_project_count", "n/a"))
        row[3].text = _fmt(project.get("scaffold_rmse"))
        row[4].text = _fmt(project.get("mean_feasibility_score"))
        row[5].text = _fmt(project.get("best_vina_affinity_kcal"))


def _add_project_snapshot(document: Document, snapshot: dict) -> None:
    project = snapshot.get("project_snapshot", {})
    document.add_heading("Current OncoForge Snapshot", level=1)
    document.add_paragraph(
        "This section keeps the glossary tied to the live project state so the document stays useful as the code evolves."
    )
    _add_label(document, "Dataset", str(project.get("dataset_name", "n/a")))
    _add_label(document, "Scaffold RMSE", _fmt(project.get("scaffold_rmse")))
    _add_label(document, "Audit pass rate", _fmt(project.get("audit_pass_rate")))
    _add_label(document, "Mean feasibility", _fmt(project.get("mean_feasibility_score")))
    _add_label(document, "Cross-database mean consensus", _fmt(project.get("cross_database_mean_consensus")))
    _add_label(document, "Cross-database strong rate", _fmt(project.get("cross_database_strong_rate")))
    _add_label(document, "External evidence mean support", _fmt(project.get("mean_external_evidence_support")))
    _add_label(document, "Source holdout mean RMSE", _fmt(project.get("source_holdout_mean_rmse")))
    _add_label(document, "Rediscovery protected top-10 recall", _fmt(project.get("rediscovery_protected_top10_recall")))
    _add_label(document, "RL external evidence mean support", _fmt(project.get("rl_mean_external_evidence_support")))
    _add_label(document, "PubChem mean enriched evidence", _fmt(project.get("pubchem_mean_enriched_evidence_score")))
    _add_label(document, "Best Vina affinity", _fmt(project.get("best_vina_affinity_kcal")))
    _add_label(document, "Mean interaction support", _fmt(project.get("mean_interaction_support")))
    _add_label(document, "RL best episode return", _fmt(project.get("rl_best_episode_return")))
    rl_top = project.get("rl_top_candidate") or {}
    if rl_top:
        _add_label(
            document,
            "Current RL lead",
            f"pIC50={_fmt(rl_top.get('predicted_pIC50'))}, "
            f"docking={_fmt(rl_top.get('docking_rescore'))}, "
            f"interaction={_fmt(rl_top.get('interaction_support_score'))}, "
            f"feasibility={_fmt(rl_top.get('feasibility_score'))}",
        )


def _add_category_summary(document: Document) -> None:
    document.add_heading("Category Summary", level=1)
    table = document.add_table(rows=1, cols=3)
    header = table.rows[0].cells
    header[0].text = "Category"
    header[1].text = "Terms"
    header[2].text = "Used Directly"
    for category in CATEGORY_ORDER:
        entries = [entry for entry in BUZZWORD_ENTRIES if entry.category == category]
        if not entries:
            continue
        row = table.add_row().cells
        row[0].text = category
        row[1].text = str(len(entries))
        row[2].text = str(sum(1 for entry in entries if entry.used_in_project))


def _add_term(document: Document, entry: BuzzwordEntry) -> None:
    document.add_heading(entry.term, level=2)
    _add_label(document, "Used in OncoForge now", "Yes" if entry.used_in_project else "Not yet; included as core context")
    _add_label(document, "Simple explanation", entry.short_definition)
    _add_label(document, "Deeper explanation", entry.detailed_explanation)
    _add_label(document, "Why it matters", entry.why_it_matters)
    _add_label(document, "How OncoForge uses it", entry.oncoforge_usage)
    _add_label(document, "Common pitfall", entry.common_pitfall)
    if entry.related_terms:
        _add_label(document, "Related terms", ", ".join(entry.related_terms))


def _build_docx(snapshot: dict, history: list[dict]) -> Path:
    document = Document()
    _add_title(
        document,
        "OncoForge Buzzword Glossary: Machine Learning, Chemistry, Biology, and Project Concepts",
    )
    document.add_paragraph(
        "This is a living glossary for the project. It explains the core language used in OncoForge and in modern AI-driven molecular discovery in a way that is easier to read than a typical paper."
    )
    document.add_paragraph(
        "Each term includes a plain-language definition, a deeper explanation, why it matters, how it appears inside OncoForge, and one common misunderstanding to avoid."
    )
    _add_history_table(document, history)
    _add_project_snapshot(document, snapshot)
    _add_category_summary(document)

    for category in CATEGORY_ORDER:
        entries = sorted((entry for entry in BUZZWORD_ENTRIES if entry.category == category), key=lambda entry: entry.term.lower())
        if not entries:
            continue
        document.add_section(WD_SECTION.NEW_PAGE)
        document.add_heading(category, level=1)
        for entry in entries:
            _add_term(document, entry)

    document.add_section(WD_SECTION.NEW_PAGE)
    document.add_heading("How To Update This Document", level=1)
    document.add_paragraph(
        "The source of truth for glossary entries is src/knowledge/oncoforge_buzzwords.py. "
        "Whenever new concepts become important in the codebase, update that file and rerun the builder."
    )
    document.add_paragraph("Builder command: python -m src.pipelines.build_buzzword_glossary_docx")

    DOCX_PATH.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(DOCX_PATH))
    return DOCX_PATH


def main(argv: list[str] | None = None) -> None:
    history = _load_history()
    snapshot = _prepare_snapshot(history)
    history = [entry for entry in history if entry.get("run_label") != snapshot.get("run_label")]
    history.append(snapshot)
    history.sort(key=lambda entry: entry.get("created_at", ""))
    _save_history(history)
    _append_context_memory(snapshot)
    out_path = _build_docx(snapshot, history)
    print(f"[OK] Saved buzzword glossary Word document: {out_path}")
    print(f"[OK] Updated glossary history: {HISTORY_INDEX}")


if __name__ == "__main__":
    main()
