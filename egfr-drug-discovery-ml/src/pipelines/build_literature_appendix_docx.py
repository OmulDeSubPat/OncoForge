from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH

from src.config import PROJECT_ROOT
from src.knowledge import COMPETITION_LITERATURE, PROJECT_PHASES


REPORTS_DIR = PROJECT_ROOT / "reports"
DOCX_PATH = REPORTS_DIR / "OncoForge_Literature_Appendix.docx"


def _add_title(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = document.styles["Title"].font.size


def _add_label(document: Document, label: str, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.add_run(f"{label}: ").bold = True
    paragraph.add_run(text)


def build_docx(path: Path | None = None) -> Path:
    out_path = path or DOCX_PATH
    document = Document()

    _add_title(document, "OncoForge Literature Appendix and External Comparison Notes")
    document.add_paragraph(
        "This appendix accompanies the technical notebook and glossary. "
        "It groups the main references used to position OncoForge against foundational cheminformatics work, "
        "public bioactivity databases, generative molecular design systems, and recent EGFR-specific AI studies."
    )
    document.add_paragraph(
        "The short quotes are intentionally brief. External comparison values are context markers, not strict leaderboard claims, "
        "because the referenced studies differ in endpoint, dataset scope, and validation protocol."
    )

    document.add_heading("Project Evolution Reference Frame", level=1)
    for phase in PROJECT_PHASES:
        document.add_heading(f"{phase.phase_id} | {phase.title}", level=2)
        _add_label(document, "Date / commit", f"{phase.date_label} / {phase.commit}")
        _add_label(document, "Focus", phase.focus)
        for upgrade in phase.upgrades:
            document.add_paragraph(upgrade, style="List Bullet")

    categories = sorted({entry.category for entry in COMPETITION_LITERATURE})
    for category in categories:
        document.add_section(WD_SECTION.NEW_PAGE)
        document.add_heading(category, level=1)
        category_entries = [entry for entry in COMPETITION_LITERATURE if entry.category == category]
        for entry in category_entries:
            document.add_heading(entry.title, level=2)
            _add_label(document, "Citation", entry.citation)
            _add_label(document, "Why it matters for OncoForge", entry.why_it_matters)
            _add_label(document, "Short quote", f"\"{entry.short_quote}\"")
            _add_label(document, "Source", entry.url)
            if entry.comparison_label and entry.comparison_value is not None:
                unit = f" {entry.comparison_unit}" if entry.comparison_unit else ""
                _add_label(document, "Comparison signal", f"{entry.comparison_label}: {entry.comparison_value}{unit}")
            if entry.comparison_note:
                _add_label(document, "Comparison note", entry.comparison_note)

    document.add_section(WD_SECTION.NEW_PAGE)
    document.add_heading("Reference Count", level=1)
    document.add_paragraph(f"Total literature entries included: {len(COMPETITION_LITERATURE)}")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(out_path))
    return out_path


def main() -> None:
    out_path = build_docx()
    print(f"[OK] Saved literature appendix Word document: {out_path}")


if __name__ == "__main__":
    main()
