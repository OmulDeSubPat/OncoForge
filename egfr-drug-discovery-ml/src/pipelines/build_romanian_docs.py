from __future__ import annotations

import json
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

from src.config import PROJECT_ROOT
from src.knowledge import BUZZWORD_ENTRIES, COMPETITION_LITERATURE, PROJECT_PHASES


REPORTS_DIR = PROJECT_ROOT / "reports"
NOTEBOOK_DIR = REPORTS_DIR / "technical_notebook"
MULTI_AGENT_ABLATION_RO = REPORTS_DIR / "multi_agent_ablation.csv"

CAIET_RO = REPORTS_DIR / "Caiet_Tehnic_OncoForge_ISEF_RO.docx"
GLOSAR_RO = REPORTS_DIR / "Glosar_Tehnic_OncoForge_RO.docx"
LITERATURA_RO = REPORTS_DIR / "Anexa_Literatura_OncoForge_RO.docx"


def _load_json(path: Path) -> dict | list | None:
    if not path.exists():
        return None
    return json.loads(path.read_text(encoding="utf-8"))


def _load_csv(path: Path) -> pd.DataFrame | None:
    return pd.read_csv(path, low_memory=False) if path.exists() else None


def _fmt(value) -> str:
    try:
        return f"{float(value):.3f}"
    except (TypeError, ValueError):
        return "n/a"


def _add_title(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = document.styles["Title"].font.size


def _add_label(document: Document, label: str, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.add_run(f"{label}: ").bold = True
    paragraph.add_run(text)


def _add_picture(document: Document, path: Path, width: float = 6.5) -> None:
    if path.exists():
        document.add_picture(str(path), width=Inches(width))


def _format_table(df: pd.DataFrame | None, columns: list[str], n: int = 8) -> str:
    if df is None or df.empty:
        return "Nu exista date disponibile."
    cols = [column for column in columns if column in df.columns]
    if not cols:
        return "Coloanele asteptate nu sunt disponibile."
    subset = df[cols].head(n).copy()
    header = "| " + " | ".join(cols) + " |"
    separator = "| " + " | ".join(["---"] * len(cols)) + " |"
    rows = [header, separator]
    for _, row in subset.iterrows():
        values = []
        for column in cols:
            value = row[column]
            if isinstance(value, float):
                values.append(f"{value:.3f}")
            else:
                values.append(str(value))
        rows.append("| " + " | ".join(values) + " |")
    return "\n".join(rows)


def _add_bullet(document: Document, text: str) -> None:
    document.add_paragraph(text, style="List Bullet")


def _multi_agent_ablation_summary() -> dict[str, float] | None:
    ablation_df = _load_csv(MULTI_AGENT_ABLATION_RO)
    if ablation_df is None or ablation_df.empty:
        return None
    subset = ablation_df[ablation_df["strategy"].isin(["protected_final", "naive_proxy"])].copy()
    if subset.empty:
        return None
    top_k_values = sorted(set(subset["top_k"].tolist()))
    if not top_k_values:
        return None
    selected_top_k = 100 if 100 in top_k_values else top_k_values[min(len(top_k_values) - 1, 1)]
    protected = subset[(subset["strategy"] == "protected_final") & (subset["top_k"] == selected_top_k)]
    naive = subset[(subset["strategy"] == "naive_proxy") & (subset["top_k"] == selected_top_k)]
    if protected.empty or naive.empty:
        return None
    protected_row = protected.iloc[0]
    naive_row = naive.iloc[0]
    return {
        "top_k": float(selected_top_k),
        "protected_pic50": float(protected_row.get("mean_predicted_pIC50", 0.0)),
        "naive_pic50": float(naive_row.get("mean_predicted_pIC50", 0.0)),
        "protected_risk": float(protected_row.get("mean_reward_hacking_risk", 0.0)),
        "naive_risk": float(naive_row.get("mean_reward_hacking_risk", 0.0)),
        "protected_pass": float(protected_row.get("audit_pass_rate", 0.0)),
        "naive_pass": float(naive_row.get("audit_pass_rate", 0.0)),
    }


def _figure_guides_ro(model_summary: dict, metrics: dict) -> dict[str, dict[str, str]]:
    multi_agent = _multi_agent_ablation_summary()
    if multi_agent:
        multi_agent_text = (
            f"Figura compara selectorul simplu, bazat aproape doar pe scorul de recompensa, cu sistemul multi-agent protejat. "
            f"La top-{int(multi_agent['top_k'])}, varianta protejata pastreaza un pIC50 mediu de {_fmt(multi_agent['protected_pic50'])}, "
            f"foarte aproape de selectorul simplu ({_fmt(multi_agent['naive_pic50'])}), dar scade riscul mediu de reward hacking de la "
            f"{_fmt(multi_agent['naive_risk'])} la {_fmt(multi_agent['protected_risk'])}."
        )
        multi_agent_importance = (
            f"Este importanta deoarece arata exact avantajul arhitecturii proiectului: audit pass rate creste de la {_fmt(multi_agent['naive_pass'])} "
            f"la {_fmt(multi_agent['protected_pass'])}, deci sistemul nu doar sorteaza altfel moleculele, ci le face mai credibile pentru selectie."
        )
    else:
        multi_agent_text = "Figura compara selectorul simplu bazat pe scor cu sistemul multi-agent protejat pe shortlist-uri de dimensiuni diferite."
        multi_agent_importance = "Este importanta deoarece arata daca stratul de audit si consens chiar imbunatateste calitatea shortlist-ului."

    return {
        "pipeline_flowchart.png": {
            "titlu": "Schema simplificata a pipeline-ului",
            "ce_arata": "Figura rezuma traseul complet al proiectului: de la datele publice EGFR si curatarea lor, pana la modelul QSAR, generarea de candidati, rankingul multi-agent, validarea structurala si batch-ul prospectiv.",
            "importanta": "Este importanta fiindca ajuta juriul sa vada ca proiectul nu este un singur model, ci un lant de verificari succesive care elimina riscuri diferite.",
        },
        "project_phase_capability_matrix.png": {
            "titlu": "Matricea capabilitatilor pe iteratii",
            "ce_arata": "Figura arata ce module au fost adaugate in fiecare faza: validare pe mai multe split-uri, audit, scorare structurala, verificare cross-database si selectie prospectiva.",
            "importanta": "Este importanta deoarece demonstreaza evolutia reala a proiectului. Fiecare iteratie a adaugat control si validare, nu doar complexitate.",
        },
        "project_evolution_history.png": {
            "titlu": "Cronologia evolutiei proiectului",
            "ce_arata": "Figura urmareste trecerea de la versiunea de baza de pe Desktop la platforma actuala, orientata spre prioritizare auditabila si suport extern.",
            "importanta": "Este importanta pentru competitie deoarece pune in evidenta procesul de cercetare si modul in care proiectul s-a maturizat prin iteratii.",
        },
        "single_agent_vs_multi_agent.png": {
            "titlu": "Single-agent vs multi-agent",
            "ce_arata": multi_agent_text,
            "importanta": multi_agent_importance,
        },
        "naive_vs_protected_scores.png": {
            "titlu": "Scor naiv vs scor protejat",
            "ce_arata": "Figura compara scorul brut cu scorul final protejat si evidentiaza moleculele care erau supraevaluate inainte de audit.",
            "importanta": "Este importanta pentru ca arata cum se schimba shortlist-ul atunci cand proiectul nu mai are incredere doar in un singur proxy numeric.",
        },
        "model_split_performance.png": {
            "titlu": "Performanta modelului pe split-uri diferite",
            "ce_arata": (
                f"Figura compara eroarea modelului pe split aleator, scaffold si temporal. In snapshot-ul curent, RMSE este "
                f"{_fmt(model_summary.get('random_split', {}).get('rmse'))} pe random, "
                f"{_fmt(model_summary.get('scaffold_split', {}).get('rmse'))} pe scaffold si "
                f"{_fmt(model_summary.get('temporal_split', {}).get('rmse'))} pe temporal."
            ),
            "importanta": "Este importanta deoarece scaffold si temporal split sunt mult mai aproape de scenariul real de generalizare decat un split aleator simplu.",
        },
        "cross_database_consensus_vs_readiness.png": {
            "titlu": "Consens cross-database vs experimental readiness",
            "ce_arata": "Figura pune fata in fata suportul din baze de date independente si scorul de experimental readiness al candidatilor finali.",
            "importanta": f"Este importanta pentru ca arata daca moleculele trimise mai departe sunt sustinute si extern, nu doar de modelul principal. Consensul mediu actual este { _fmt(metrics.get('cross_database_mean_consensus')) }.",
        },
        "source_holdout_rmse.png": {
            "titlu": "Eroare pe benchmark-ul source holdout",
            "ce_arata": "Figura masoara eroarea atunci cand o sursa intreaga de date este scoasa din antrenare si folosita doar pentru testare.",
            "importanta": f"Este importanta deoarece testeaza transferul intre surse publice diferite. RMSE-ul mediu actual este {_fmt(metrics.get('source_holdout_mean_rmse'))}.",
        },
        "rediscovery_recall_at_k.png": {
            "titlu": "Rediscovery recall la k",
            "ce_arata": "Figura arata cat de des recupereaza rankingul protejat molecule puternice deja cunoscute in primele pozitii ale listei.",
            "importanta": f"Este importanta deoarece verifica daca sistemul stie sa recunoasca chimie valoroasa, nu doar sa inventeze structuri. Recall-ul protejat top-10 este {_fmt(metrics.get('rediscovery_protected_top10_recall'))}.",
        },
        "marketed_vs_generated_boxplots.png": {
            "titlu": "Molecule de piata vs molecule generate",
            "ce_arata": "Figura compara moleculele generate cu molecule EGFR consacrate pe axe precum potenta estimata, QED si risc de reward hacking.",
            "importanta": "Este importanta fiindca ofera un reper intuitiv pentru juriu: unde se afla candidatii noi fata de chimia deja validata in acest spatiu terapeutic.",
        },
        "technical_notebook_chemical_space.png": {
            "titlu": "Snapshot de spatiu chimic",
            "ce_arata": "Figura proiecteaza in doua dimensiuni moleculele de piata, lead-urile clasate si shortlist-ul novel pentru a arata asemanari si diferente de spatiu chimic.",
            "importanta": "Este importanta deoarece arata echilibrul dintre noutate si realism. Proiectul trebuie sa exploreze dincolo de medicamentele cunoscute fara sa iasa complet din zona sustinuta de date.",
        },
        "literature_context_comparison.png": {
            "titlu": "Comparatie contextuala cu literatura",
            "ce_arata": "Figura pozitioneaza proiectul fata de alte studii EGFR si repere din literatura, tinand cont de diferentele de protocol si de tipul de split.",
            "importanta": "Este importanta pentru ca ancoreaza rezultatele intr-un context extern si arata ca performanta obtinuta este credibila, nu izolata.",
        },
        "prospective_batch_readiness_vs_novelty.png": {
            "titlu": "Batch prospectiv: readiness vs novelty",
            "ce_arata": "Figura arata cum este echilibrat batch-ul final intre molecule mai noi si molecule mai pregatite pentru validare.",
            "importanta": "Este importanta deoarece scopul final al proiectului este un portofoliu rational pentru testare, nu doar un top bazat pe un singur scor.",
        },
    }


def _add_figure_note_ro(document: Document, plot_name: str, figure_guides: dict[str, dict[str, str]], width: float = 6.6) -> None:
    guide = figure_guides.get(plot_name)
    if not guide:
        return
    document.add_heading(guide["titlu"], level=2)
    _add_picture(document, NOTEBOOK_DIR / plot_name, width=width)
    _add_label(document, "Ce arata figura", guide["ce_arata"])
    _add_label(document, "De ce este importanta", guide["importanta"])


CATEGORY_RO = {
    "Project Architecture": "Arhitectura proiectului",
    "Machine Learning": "Machine learning",
    "Reinforcement Learning": "Reinforcement learning",
    "Cheminformatics": "Cheminformatica",
    "Medicinal Chemistry": "Chimie medicinala",
    "Structural Biology": "Biologie structurala",
    "Cancer Biology": "Biologia cancerului",
    "Drug-likeness and prioritization": "Drug-likeness si prioritizare",
    "Medicinal chemistry realism": "Realism de chimie medicinala",
    "Assay artifacts and safety filters": "Artefacte experimentale si filtre de siguranta",
    "Structure-based scoring": "Scorare bazata pe structura",
    "Reference databases": "Baze de date de referinta",
    "Generative molecular optimization": "Optimizare moleculara generativa",
    "EGFR study comparisons": "Comparatii cu studii EGFR",
}


TERM_SHORT_RO = {
    "Multi-agent system": "Un sistem in care mai multe componente specializate evalueaza aceeasi molecula din perspective diferite.",
    "Verifiable reward": "Un scor de recompensa construit din componente care pot fi urmarite si verificate dupa rulare.",
    "Reward hacking": "Situatia in care un model obtine scoruri mari fara sa rezolve cu adevarat problema stiintifica.",
    "Audit agent": "Componenta care verifica daca o molecula pare buna din motive corecte, nu doar pentru ca exploateaza modelul.",
    "Verified reward": "Scorul final protejat, dupa aplicarea penalizarilor si verificarilor de suport.",
    "Naive reward": "Scorul brut, calculat inainte de filtrele mai stricte de protectie.",
    "Applicability domain": "Zona din spatiul chimic in care modelul are suficiente exemple similare ca sa fie credibil.",
    "Novelty": "Gradul in care o molecula este diferita de cele din antrenare sau de medicamentele cunoscute.",
    "Diversity": "Masura care impiedica shortlist-ul final sa fie format din multe molecule aproape identice.",
    "Feature engineering": "Transformarea moleculelor in reprezentari numerice utile pentru model.",
    "Descriptor": "O proprietate numerica a moleculei, precum masa moleculara sau suprafata polara.",
    "ECFP / Morgan fingerprint": "O reprezentare circulara a moleculei sub forma de vector fix, utila in QSAR si similaritate.",
    "Ensemble learning": "Combinarea mai multor modele pentru predictii mai stabile.",
    "Random Forest": "Un ansamblu de arbori de decizie antrenati pe subseturi ale datelor si caracteristicilor.",
    "Extra Trees": "Un model de tip arbori asemanator cu Random Forest, dar cu mai multa aleatoritate la impartiri.",
    "HistGradientBoosting": "Un model boosting care construieste arborii secvential pentru a corecta erorile anterioare.",
    "Uncertainty estimation": "Estimarea nivelului de incredere al modelului intr-o predictie.",
    "Calibration": "Ajustarea incertitudinii astfel incat sa reflecte mai bine eroarea reala.",
    "Scaffold split": "O impartire train-test in care moleculele sunt separate dupa schelet, nu aleator.",
    "Temporal split": "O impartire a datelor dupa timp, in care trecutul antreneaza si viitorul testeaza.",
    "Source holdout benchmark": "Un test in care o sursa publica de date este lasata complet in afara antrenarii si folosita doar pentru testare.",
    "Rediscovery benchmark": "Un benchmark care verifica daca sistemul recupereaza molecule cunoscute puternice intr-un panou dificil.",
    "QSAR": "Predictia activitatii biologice pornind de la structura moleculara.",
    "Reinforcement learning": "Un cadru de invatare in care un agent ia actiuni si primeste recompense.",
    "Q-learning": "Un algoritm RL care invata cat de buna este fiecare actiune intr-o anumita stare.",
    "Replay buffer": "Memorie de experiente anterioare folosita pentru un antrenament RL mai stabil.",
    "Reward shaping": "Definirea componentelor de recompensa astfel incat agentul sa invete mai rapid si mai sigur.",
    "SMILES": "Notatie text care descrie o molecula ca sir de caractere.",
    "Canonical SMILES": "Forma standardizata de SMILES folosita pentru a evita dublurile.",
    "Scaffold / Murcko scaffold": "Scheletul molecular central ramas dupa eliminarea multor lanturi laterale.",
    "Lipinski Rule of Five": "Set de reguli empirice folosite pentru a evalua daca o molecula are proprietati compatibile cu administrarea orala.",
    "QED": "Scor compozit de drug-likeness care rezuma mai multe proprietati favorabile.",
    "Synthetic accessibility": "Estimarea cat de usor poate fi sintetizata o molecula.",
    "PAINS filter": "Filtru pentru motive structurale care pot produce semnale false in bioeseuri.",
    "Docking": "Simulare computationala a legarii unei molecule de o proteina tinta.",
    "Binding pose": "Pozitia si orientarea moleculei in situsul de legare.",
    "EGFR": "Receptorul pentru factorul de crestere epidermal, o tinta majora in oncologie.",
    "Tyrosine kinase": "Enzima care transfera grupari fosfat pe tirozina si controleaza semnalizarea celulara.",
    "Lead optimization": "Etapa in care moleculele promitatoare sunt imbunatatite pentru a deveni candidati mai credibili.",
    "Experimental readiness": "Scor care estimeaza cat de aproape este un candidat de o selectie rationala pentru validare.",
    "Prospective validation batch": "Setul final de candidati selectati pentru o evaluare viitoare, nu doar top-N dupa un scor.",
    "Acquisition score": "Scor folosit pentru a echilibra exploatarea, noutatea, incertitudinea si diversitatea.",
    "Papyrus": "Baza de date publica, curatata la scara mare, folosita pentru predictii de bioactivitate.",
    "ExCAPE-DB": "Baza de date chemogenomica de mari dimensiuni, utila ca dovada externa suplimentara.",
    "IUPHAR / Guide to Pharmacology": "Resursa expert-curatata cu relatii tinta-ligand si date farmacologice.",
    "PubChem BioAssay": "Colectie mare de bioeseuri, utila ca sursa independenta de dovezi pentru EGFR.",
}


def _category_ro(category: str) -> str:
    return CATEGORY_RO.get(category, category)


def _why_ro(entry) -> str:
    if entry.used_in_project:
        return "Conteaza direct in OncoForge deoarece influenteaza calitatea shortlist-ului, credibilitatea predictiilor sau validarea externa."
    if entry.category == "Cancer Biology":
        return "Conteaza pentru baza biologica a proiectului si pentru justificarea alegerii tintei EGFR."
    if entry.category == "Structural Biology":
        return "Conteaza pentru intelegerea modului in care suportul structural completeaza predictiile QSAR."
    return "Conteaza pentru contextul teoretic al proiectului si pentru explicarea limbajului tehnic folosit in raport."


def _usage_ro(entry) -> str:
    if entry.term in TERM_SHORT_RO:
        if entry.used_in_project:
            return "Termenul este folosit direct in pipeline-ul OncoForge, in ranking, selectie, validare sau interpretarea candidatilor."
        return "Termenul ofera context util pentru intelegerea proiectului, chiar daca nu este un modul central al pipeline-ului curent."
    if entry.used_in_project:
        return "Apare direct in evaluarea sau explicarea rezultatelor OncoForge."
    return "Este inclus pentru context teoretic si pentru prezentarea mai clara a proiectului."


def _pitfall_ro(entry) -> str:
    if entry.term == "Reward hacking":
        return "Capcana principala este sa confunzi un scor mare cu un candidat bun, fara sa verifici dovezile independente."
    if entry.term == "Scaffold split":
        return "Capcana principala este sa raportezi doar split aleator si sa supraestimezi performanta reala."
    if entry.term == "Docking":
        return "Capcana principala este sa tratezi docking-ul ca dovada experimentala, desi el ofera doar suport computational."
    return "Capcana principala este interpretarea prea optimista a unui indicator fara validare suplimentara."


def _build_caiet_tehnic_ro() -> Path:
    model_summary = _load_json(REPORTS_DIR / "model_performance_summary.json") or {}
    metrics = _load_json(NOTEBOOK_DIR / "technical_notebook_metrics.json") or {}
    context = _load_json(NOTEBOOK_DIR / "competition_report_context.json") or {}
    history = _load_json(REPORTS_DIR / "technical_notebook_history" / "run_history.json") or []

    prospective_df = _load_csv(REPORTS_DIR / "prospective_validation_batch.csv")
    marketed_df = _load_csv(REPORTS_DIR / "marketed_egfr_scored.csv")
    shortlist_df = _load_csv(REPORTS_DIR / "market_comparable_novel_shortlist.csv")
    diverse_df = _load_csv(REPORTS_DIR / "final_diverse_candidates.csv")
    crossdb_df = _load_csv(REPORTS_DIR / "iterative_ai_optimized_candidates_structural_crossdb.csv")

    document = Document()
    _add_title(document, "Caiet Tehnic OncoForge in Limba Romana")
    document.add_paragraph(
        "Acest document este versiunea in limba romana a caietului tehnic OncoForge. "
        "El rezuma fundamentele biologice si chimice ale proiectului, metodologia de machine learning, evolutia iteratiilor, "
        "rezultatele cheie si pozitionarea fata de literatura de specialitate."
    )

    document.add_heading("Rezumat", level=1)
    document.add_paragraph(
        "OncoForge este o platforma computationala pentru optimizarea lead-urilor EGFR. "
        "Scopul ei nu este sa declare descoperirea unui medicament final, ci sa prioritizeze candidati moleculari mai credibili pentru validari ulterioare. "
        "Sistemul combina predictia QSAR, scorare multi-agent, audit anti-reward-hacking, rescoring structural, validare cross-database, "
        "evaluare de fezabilitate si experimente de selectie prospectiva."
    )
    document.add_paragraph(
        f"In configuratia actuala, modelul foloseste {int(context.get('model_dataset_size', 0) or 0)} molecule curate pentru antrenare, "
        f"raporteaza RMSE {_fmt(model_summary.get('random_split', {}).get('rmse'))} pe split aleator si RMSE {_fmt(model_summary.get('scaffold_split', {}).get('rmse'))} pe scaffold split, "
        f"iar shortlist-ul prospectiv are {int(metrics.get('prospective_batch_size', 0) or 0)} candidati."
    )

    document.add_heading("Fundament Biologic si Chimic", level=1)
    document.add_paragraph(
        "EGFR este o tinta relevanta in oncologie deoarece dereglarea caii sale de semnalizare poate sustine proliferarea necontrolata a celulelor. "
        "De aceea, inhibarea domeniului tirozin-kinazic EGFR este o strategie importanta in dezvoltarea de molecule anticancer."
    )
    document.add_paragraph(
        "Pe partea de chimie medicinala, problema nu este doar sa generezi molecule noi, ci sa alegi molecule care sunt plauzibile, suficient de active, "
        "compatibile cu proprietati de tip drug-likeness, fezabile sintetic si sustinute de dovezi externe sau structurale."
    )

    document.add_heading("Metodologie de Machine Learning", level=1)
    document.add_paragraph(
        "Componenta principala este un ansamblu multiview pentru predictia potenței EGFR, evaluat pe split aleator, scaffold si temporal. "
        "Dupa predictia initiala, moleculele sunt reordonate de un sistem multi-agent care separa potenta, chimia, siguranta, noutatea, aplicabilitatea si auditul anti-reward-hacking."
    )
    document.add_paragraph(
        "Proiectul evita sa foloseasca un singur scor brut. In loc de asta, compara un scor naiv cu un scor protejat si penalizeaza candidatii care ies din zona de incredere a modelului sau par sa exploateze proxy-uri fragile."
    )

    document.add_heading("Evolutia Proiectului", level=1)
    for phase in PROJECT_PHASES:
        document.add_heading(f"{phase.phase_id} | {phase.title}", level=2)
        _add_label(document, "Data / commit", f"{phase.date_label} / {phase.commit}")
        _add_label(document, "Focus", phase.focus)
        for upgrade in phase.upgrades:
            document.add_paragraph(upgrade, style="List Bullet")

    _add_picture(document, NOTEBOOK_DIR / "project_phase_capability_matrix.png", width=6.8)
    _add_picture(document, NOTEBOOK_DIR / "project_evolution_history.png", width=6.8)

    document.add_heading("Rezultate Cheie", level=1)
    _add_label(document, "Random split RMSE / R2", f"{_fmt(model_summary.get('random_split', {}).get('rmse'))} / {_fmt(model_summary.get('random_split', {}).get('r2'))}")
    _add_label(document, "Scaffold split RMSE / R2", f"{_fmt(model_summary.get('scaffold_split', {}).get('rmse'))} / {_fmt(model_summary.get('scaffold_split', {}).get('r2'))}")
    _add_label(document, "Temporal split RMSE / R2", f"{_fmt(model_summary.get('temporal_split', {}).get('rmse'))} / {_fmt(model_summary.get('temporal_split', {}).get('r2'))}")
    _add_label(document, "Rata de trecere audit", _fmt(metrics.get("audit_pass_rate")))
    _add_label(document, "Fezabilitate medie", _fmt(metrics.get("mean_feasibility_score")))
    _add_label(document, "Consens cross-database mediu", _fmt(metrics.get("cross_database_mean_consensus")))
    _add_label(document, "Cea mai buna afinitate Vina", f"{_fmt(metrics.get('best_vina_affinity_kcal'))} kcal/mol")
    _add_label(document, "Dimensiune batch prospectiv", str(int(metrics.get("prospective_batch_size", 0) or 0)))
    _add_label(document, "Rediscovery top-10 protejat", _fmt(metrics.get("rediscovery_protected_top10_recall")))

    document.add_heading("Figuri Recomandate pentru Prezentare", level=1)
    for plot_name in [
        "literature_context_comparison.png",
        "naive_vs_protected_scores.png",
        "cross_database_consensus_vs_readiness.png",
        "rediscovery_recall_at_k.png",
        "source_holdout_rmse.png",
        "marketed_vs_generated_boxplots.png",
        "technical_notebook_chemical_space.png",
    ]:
        _add_picture(document, NOTEBOOK_DIR / plot_name, width=6.6)

    document.add_heading("Candidate Reprezentativi", level=1)
    if prospective_df is not None and not prospective_df.empty:
        document.add_paragraph("Mai jos este shortlist-ul prospectiv, adica setul cel mai apropiat de o selectie rationala pentru pasi urmatori.")
        document.add_paragraph(_format_table(
            prospective_df,
            [
                "prospective_batch_rank",
                "candidate_source",
                "predicted_pIC50",
                "experimental_readiness_score",
                "prospective_acquisition_score",
                "experimental_readiness_status",
            ],
            n=10,
        ))
    if diverse_df is not None and not diverse_df.empty:
        document.add_heading("Candidati diversi", level=2)
        document.add_paragraph(_format_table(
            diverse_df,
            ["smiles", "predicted_pIC50", "QED", "reward_hacking_risk", "final_score"],
            n=8,
        ))
    if marketed_df is not None and not marketed_df.empty:
        document.add_heading("Comparație cu molecule EGFR comercializate", level=2)
        document.add_paragraph(_format_table(
            marketed_df,
            ["name", "predicted_pIC50", "vina_affinity_kcal", "interaction_support_score", "final_score"],
            n=8,
        ))
    if shortlist_df is not None and not shortlist_df.empty:
        document.add_heading("Shortlist novel comparabil cu piata", level=2)
        document.add_paragraph(_format_table(
            shortlist_df,
            ["smiles", "predicted_pIC50", "QED", "max_market_similarity", "final_score"],
            n=8,
        ))

    document.add_heading("Validare Externa si Benchmarkuri", level=1)
    document.add_paragraph(
        "Punctul forte al proiectului este validarea in afara scorului principal. "
        "Pe langa split-urile standard, OncoForge foloseste source holdout benchmark, rediscovery benchmark, challenge anti-reward-hacking si dovezi independente din baze de date precum Papyrus, ExCAPE-DB, PubChem si IUPHAR."
    )
    if crossdb_df is not None and not crossdb_df.empty:
        document.add_paragraph(_format_table(
            crossdb_df,
            [
                "smiles",
                "predicted_pIC50",
                "cross_database_consensus_score",
                "external_evidence_support",
                "cross_database_status",
            ],
            n=8,
        ))

    document.add_heading("Comparatie cu Literatura", level=1)
    document.add_paragraph(
        "Comparațiile cu literatura trebuie interpretate ca repere contextuale, nu ca un clasament perfect, deoarece studiile folosesc seturi de date, endpoint-uri si protocoale diferite."
    )
    _add_picture(document, NOTEBOOK_DIR / "literature_context_comparison.png", width=6.8)
    for entry in COMPETITION_LITERATURE[:10]:
        document.add_heading(entry.title, level=2)
        _add_label(document, "Citare", entry.citation)
        _add_label(document, "De ce conteaza aici", entry.why_it_matters)
        _add_label(document, "Citat scurt", f"\"{entry.short_quote}\"")
        _add_label(document, "Sursa", entry.url)

    document.add_heading("Limitari", level=1)
    document.add_paragraph(
        "Acesta este un proiect de prioritizare computationala, nu dovada unui medicament final. "
        "Docking-ul si scorurile QSAR sunt suport computational, nu validare experimentala. "
        "Split-ul temporal ramane dificil, iar ramurile neurale GPU sunt mai degraba extensii de cercetare decat sursa principala a celor mai puternice concluzii."
    )

    document.add_heading("Concluzie", level=1)
    document.add_paragraph(
        "Contributia centrala a proiectului este trecerea de la o cautare bazata pe un singur scor la un sistem auditabil, multi-criterial si sustinut de dovezi independente. "
        "Aceasta abordare este mai potrivita pentru competitie, deoarece arata nu doar performanta numerica, ci si controlul erorilor, validarea externa si maturizarea proiectului prin iteratii."
    )

    CAIET_RO.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(CAIET_RO))
    return CAIET_RO


def _build_glosar_ro() -> Path:
    metrics = _load_json(NOTEBOOK_DIR / "technical_notebook_metrics.json") or {}
    document = Document()
    _add_title(document, "Glosar Tehnic OncoForge in Limba Romana")
    document.add_paragraph(
        "Acest glosar este o versiune practica in limba romana, gandita pentru prezentare, sustinere si intelegerea rapida a conceptelor folosite in OncoForge."
    )
    document.add_paragraph(
        f"Snapshot curent: audit pass rate {_fmt(metrics.get('audit_pass_rate'))}, scaffold RMSE {_fmt(metrics.get('model_scaffold_rmse'))}, "
        f"fezabilitate medie {_fmt(metrics.get('mean_feasibility_score'))}, consens cross-database {_fmt(metrics.get('cross_database_mean_consensus'))}."
    )

    categories = []
    for entry in BUZZWORD_ENTRIES:
        if entry.category not in categories:
            categories.append(entry.category)

    for category in categories:
        document.add_section(WD_SECTION.NEW_PAGE)
        document.add_heading(_category_ro(category), level=1)
        entries = [entry for entry in BUZZWORD_ENTRIES if entry.category == category]
        for entry in entries:
            document.add_heading(entry.term, level=2)
            _add_label(document, "Folosit direct in OncoForge", "Da" if entry.used_in_project else "Mai ales ca fundal teoretic")
            _add_label(document, "Explicatie pe scurt", TERM_SHORT_RO.get(entry.term, f"Concept din categoria {_category_ro(category).lower()}, relevant pentru intelegerea proiectului si a modului in care sunt evaluate moleculele."))
            _add_label(document, "De ce conteaza", _why_ro(entry))
            _add_label(document, "Cum apare in OncoForge", _usage_ro(entry))
            _add_label(document, "Capcana frecventa", _pitfall_ro(entry))
            if entry.related_terms:
                _add_label(document, "Termeni inruditi", ", ".join(entry.related_terms))

    document.add_section(WD_SECTION.NEW_PAGE)
    document.add_heading("Cum se foloseste glosarul", level=1)
    document.add_paragraph(
        "Pentru sustinere, cei mai importanti termeni sunt: EGFR, QSAR, scaffold split, applicability domain, reward hacking, audit agent, "
        "experimental readiness, prospective validation batch, docking, Papyrus, ExCAPE-DB si cross-database validation."
    )

    GLOSAR_RO.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(GLOSAR_RO))
    return GLOSAR_RO


def _build_literatura_ro() -> Path:
    document = Document()
    _add_title(document, "Anexa de Literatura OncoForge in Limba Romana")
    document.add_paragraph(
        "Aceasta anexa reuneste literatura folosita pentru pozitionarea proiectului fata de biologia EGFR, bazele de date publice, "
        "cheminformatica, optimizarea moleculara si studiile recente de AI aplicate pe EGFR."
    )
    document.add_paragraph(
        "Titlurile lucrarilor si citatele scurte sunt pastrate in forma originala pentru fidelitate fata de sursa, dar comentariile de interpretare sunt scrise in limba romana."
    )

    document.add_heading("Cadru de evolutie al proiectului", level=1)
    for phase in PROJECT_PHASES:
        document.add_heading(f"{phase.phase_id} | {phase.title}", level=2)
        _add_label(document, "Data / commit", f"{phase.date_label} / {phase.commit}")
        _add_label(document, "Rol in evolutie", phase.focus)
        for upgrade in phase.upgrades:
            document.add_paragraph(upgrade, style="List Bullet")

    categories = []
    for entry in COMPETITION_LITERATURE:
        if entry.category not in categories:
            categories.append(entry.category)

    for category in categories:
        document.add_section(WD_SECTION.NEW_PAGE)
        document.add_heading(_category_ro(category), level=1)
        entries = [entry for entry in COMPETITION_LITERATURE if entry.category == category]
        for entry in entries:
            document.add_heading(entry.title, level=2)
            _add_label(document, "Citare", entry.citation)
            _add_label(document, "De ce este relevanta", entry.why_it_matters)
            _add_label(document, "Citat scurt", f"\"{entry.short_quote}\"")
            _add_label(document, "Sursa", entry.url)
            if entry.comparison_label and entry.comparison_value is not None:
                unit = f" {entry.comparison_unit}" if entry.comparison_unit else ""
                _add_label(document, "Semnal comparativ", f"{entry.comparison_label}: {entry.comparison_value}{unit}")
            if entry.comparison_note:
                _add_label(document, "Nota de comparatie", entry.comparison_note)

    document.add_section(WD_SECTION.NEW_PAGE)
    document.add_heading("Concluzie de pozitionare", level=1)
    document.add_paragraph(
        "Literatura sustine doua idei importante pentru proiect: in primul rand, EGFR ramane o tinta biologica relevantă; "
        "in al doilea rand, proiectele moderne de AI in chimie trebuie evaluate cu prudenta, folosind split-uri dificile, validare externa, "
        "dovezi structurale si control al reward hacking-ului."
    )
    document.add_paragraph(f"Numar total de referinte incluse: {len(COMPETITION_LITERATURE)}")

    LITERATURA_RO.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(LITERATURA_RO))
    return LITERATURA_RO


def main() -> None:
    caiet = _build_caiet_tehnic_ro()
    glosar = _build_glosar_ro()
    literatura = _build_literatura_ro()
    print(f"[OK] Saved Romanian technical notebook: {caiet}")
    print(f"[OK] Saved Romanian glossary: {glosar}")
    print(f"[OK] Saved Romanian literature appendix: {literatura}")


if __name__ == "__main__":
    main()
