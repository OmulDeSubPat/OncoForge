from __future__ import annotations

import json
from datetime import datetime
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

from src.config import PROJECT_ROOT
from src.knowledge import COMPETITION_LITERATURE, PROJECT_PHASES
from src.visualization.technical_notebook_molecule_views import (
    build_candidate_3d_views,
    build_candidate_grid,
)


REPORTS_DIR = PROJECT_ROOT / "reports"
NOTEBOOK_DIR = REPORTS_DIR / "technical_notebook"
HISTORY_DIR = REPORTS_DIR / "technical_notebook_history"
DOCX_PATH = REPORTS_DIR / "OncoForge_Technical_Notebook.docx"
ENGLISH_ALIAS_PATH = REPORTS_DIR / "Caiet_Tehnic_OncoForge_ISEF.docx"
HISTORY_INDEX = HISTORY_DIR / "run_history.json"
CONTEXT_MEMORY = HISTORY_DIR / "context_memory.md"
COMPETITION_CONTEXT_JSON = NOTEBOOK_DIR / "competition_report_context.json"
MULTI_AGENT_ABLATION = REPORTS_DIR / "multi_agent_ablation.csv"

PLOT_ORDER = [
    "pipeline_flowchart.png",
    "project_phase_capability_matrix.png",
    "project_evolution_history.png",
    "literature_context_comparison.png",
    "single_agent_vs_multi_agent.png",
    "risk_distribution_by_audit_status.png",
    "naive_vs_protected_scores.png",
    "audit_rank_demotions.png",
    "top_leads_agent_support_heatmap.png",
    "novelty_vs_applicability.png",
    "model_split_performance.png",
    "gpu_gnn_scaffold_benchmark.png",
    "uncertainty_calibration.png",
    "structural_rescoring_scatter.png",
    "vina_affinity_vs_priority.png",
    "interaction_support_vs_vina.png",
    "readiness_vs_structure.png",
    "cross_database_vs_potency.png",
    "cross_database_consensus_vs_readiness.png",
    "cross_database_status_counts.png",
    "external_evidence_support_vs_potency.png",
    "pubchem_assay_relevance.png",
    "feasibility_vs_potency.png",
    "generator_benchmark_overview.png",
    "rl_training_curve.png",
    "rl_reward_breakdown.png",
    "rl_external_evidence_vs_priority.png",
    "gpu_rl_training_curve.png",
    "gpu_actor_critic_training_curve.png",
    "model_robustness_scaffold.png",
    "source_holdout_rmse.png",
    "source_holdout_recall.png",
    "rediscovery_recall_at_k.png",
    "rediscovery_rank_shift.png",
    "challenge_rank_shift.png",
    "challenge_status_rates.png",
    "prospective_batch_readiness_vs_novelty.png",
    "marketed_vs_generated_boxplots.png",
    "structural_benchmark_boxplots.png",
    "technical_notebook_chemical_space.png",
]


def _fmt(value) -> str:
    try:
        return f"{float(value):.3f}"
    except (TypeError, ValueError):
        return "n/a"


def _load_df(path: Path) -> pd.DataFrame | None:
    return pd.read_csv(path, low_memory=False) if path.exists() else None


def _load_json(path: Path) -> dict | None:
    if not path.exists():
        return None
    return json.loads(path.read_text(encoding="utf-8"))


def _select_candidate_source() -> tuple[pd.DataFrame, str]:
    candidates = [
        (REPORTS_DIR / "prospective_validation_batch.csv", "prospective_batch"),
        (REPORTS_DIR / "iterative_ai_optimized_candidates_feasibility.csv", "optimized_feasibility"),
        (REPORTS_DIR / "final_diverse_candidates.csv", "final_diverse"),
        (REPORTS_DIR / "rl_verifiable" / "rl_top_candidates.csv", "verifiable_rl"),
        (REPORTS_DIR / "iterative_ai_optimized_candidates_structural_rescored.csv", "iterative_structural"),
        (REPORTS_DIR / "iterative_ai_optimized_candidates.csv", "iterative"),
    ]
    for path, label in candidates:
        df = _load_df(path)
        if df is not None and not df.empty:
            return df, label
    raise FileNotFoundError("No candidate artifact available for the technical notebook Word export.")


def _current_run_label() -> str:
    return datetime.now().strftime("%Y%m%d_%H%M%S")


def _candidate_summary(row: pd.Series) -> dict:
    summary = {
        "smiles": str(row.get("smiles", "")),
        "predicted_pIC50": float(row.get("predicted_pIC50", 0.0)),
        "QED": float(row.get("QED", 0.0)),
        "final_score": float(row.get("final_score", 0.0)),
    }
    if pd.notna(row.get("docking_rescore", None)):
        summary["docking_rescore"] = float(row["docking_rescore"])
    if pd.notna(row.get("vina_affinity_kcal", None)):
        summary["vina_affinity_kcal"] = float(row["vina_affinity_kcal"])
    if pd.notna(row.get("feasibility_score", None)):
        summary["feasibility_score"] = float(row["feasibility_score"])
    return summary


def _format_delta(current: float | None, previous: float | None, lower_is_better: bool = False) -> str:
    if current is None or previous is None:
        return "n/a"
    delta = float(current) - float(previous)
    if abs(delta) < 1e-9:
        return "unchanged"
    improved = delta < 0 if lower_is_better else delta > 0
    direction = "improved" if improved else "regressed"
    return f"{direction} by {abs(delta):.3f}"


def _build_evolution_note(snapshot: dict, previous_snapshot: dict | None) -> str:
    if previous_snapshot is None:
        return (
            "This run establishes a new recorded checkpoint for OncoForge. "
            "The notebook starts tracking model quality, candidate strength, structure support, and feasibility from this baseline."
        )

    current_model = snapshot.get("model_summary", {})
    previous_model = previous_snapshot.get("model_summary", {})
    current_metrics = snapshot.get("notebook_metrics", {})
    previous_metrics = previous_snapshot.get("notebook_metrics", {})
    current_top = (snapshot.get("top_candidates") or [{}])[0]
    previous_top = (previous_snapshot.get("top_candidates") or [{}])[0]

    scaffold_note = _format_delta(current_model.get("scaffold_rmse"), previous_model.get("scaffold_rmse"), lower_is_better=True)
    feasibility_note = _format_delta(current_metrics.get("mean_feasibility_score"), previous_metrics.get("mean_feasibility_score"))
    docking_note = _format_delta(current_metrics.get("best_vina_affinity_kcal"), previous_metrics.get("best_vina_affinity_kcal"), lower_is_better=True)
    top_score_note = _format_delta(current_top.get("final_score"), previous_top.get("final_score"))

    return (
        f"Compared with run {previous_snapshot.get('run_label', 'n/a')}, the scaffold validation profile {scaffold_note}, "
        f"mean feasibility {feasibility_note}, strongest structural support {docking_note}, and the top candidate priority score {top_score_note}. "
        f"This short note is kept on purpose so the document shows the evolution of the project without losing earlier iterations."
    )


def _prepare_run_snapshot(run_label: str) -> dict:
    candidate_df, candidate_source = _select_candidate_source()
    model_summary = _load_json(REPORTS_DIR / "model_performance_summary.json") or {}
    notebook_metrics = _load_json(NOTEBOOK_DIR / "technical_notebook_metrics.json") or {}

    run_dir = HISTORY_DIR / run_label
    image_dir = run_dir / "images"
    image_dir.mkdir(parents=True, exist_ok=True)

    top_candidates = candidate_df.head(12).copy().reset_index(drop=True)
    if "rank" not in top_candidates.columns:
        top_candidates["rank"] = top_candidates.index + 1

    grid_path = build_candidate_grid(top_candidates, image_dir / "top_candidates_grid.png")
    top3_assets = []
    for idx, (_, row) in enumerate(top_candidates.head(3).iterrows(), start=1):
        render_path = build_candidate_3d_views(row, image_dir / f"top3_candidate_{idx}_3d.png")
        top3_assets.append(
            {
                "rank": int(row.get("rank", idx)),
                "image": str(render_path) if render_path else None,
                "summary": _candidate_summary(row),
            }
        )

    copied_plots = []
    for plot_name in PLOT_ORDER:
        source = NOTEBOOK_DIR / plot_name
        if not source.exists():
            continue
        target = image_dir / plot_name
        target.write_bytes(source.read_bytes())
        copied_plots.append(str(target))

    rl_candidate_df = _load_df(REPORTS_DIR / "rl_verifiable" / "rl_top_candidates.csv")
    top_candidate_records = [_candidate_summary(row) for _, row in top_candidates.head(5).iterrows()]
    context_notes = [
        f"Candidate source for this run: {candidate_source}.",
        f"Top candidate score={_fmt(top_candidate_records[0].get('final_score') if top_candidate_records else None)}, "
        f"pIC50={_fmt(top_candidate_records[0].get('predicted_pIC50') if top_candidate_records else None)}, "
        f"Vina={_fmt(top_candidate_records[0].get('vina_affinity_kcal') if top_candidate_records else None)}." if top_candidate_records else "No top candidate was available.",
        f"Notebook mean feasibility={_fmt(notebook_metrics.get('mean_feasibility_score'))}, "
        f"best Vina={_fmt(notebook_metrics.get('best_vina_affinity_kcal'))}, "
        f"mean readiness={_fmt(notebook_metrics.get('mean_experimental_readiness'))}, "
        f"broad generator count={_fmt(notebook_metrics.get('generated_candidate_count'))}, "
        f"iterative generator priority={_fmt(notebook_metrics.get('iterative_mean_generator_priority'))}, "
        f"iterative parent improvement={_fmt(notebook_metrics.get('iterative_parent_improvement_rate_final_score'))}, "
        f"evidence arbiter={_fmt(notebook_metrics.get('evidence_arbiter_mean_support'))}, "
        f"cross-db mean consensus={_fmt(notebook_metrics.get('cross_database_mean_consensus'))}, "
        f"cross-db strong rate={_fmt(notebook_metrics.get('cross_database_strong_rate'))}, "
        f"external evidence support={_fmt(notebook_metrics.get('external_evidence_mean_support'))}, "
        f"RL external evidence={_fmt(notebook_metrics.get('rl_mean_external_evidence_support'))}, "
        f"GPU GNN scaffold RMSE={_fmt(notebook_metrics.get('gpu_gnn_best_scaffold_rmse'))}, "
        f"GPU RL external evidence={_fmt(notebook_metrics.get('gpu_rl_mean_external_evidence_support'))}, "
        f"GPU actor-critic external evidence={_fmt(notebook_metrics.get('gpu_actor_critic_mean_external_evidence_support'))}, "
        f"source holdout RMSE={_fmt(notebook_metrics.get('source_holdout_mean_rmse'))}, "
        f"rediscovery protected top10={_fmt(notebook_metrics.get('rediscovery_protected_top10_recall'))}, "
        f"prospective batch size={_fmt(notebook_metrics.get('prospective_batch_size'))}, "
        f"RL best episode return={_fmt(notebook_metrics.get('rl_best_episode_return'))}, "
        f"GPU actor-critic best episode return={_fmt(notebook_metrics.get('gpu_actor_critic_best_episode_return'))}.",
    ]
    history = _load_history()
    previous_snapshot = history[-1] if history else None
    evolution_note = _build_evolution_note(snapshot={
        "run_label": run_label,
        "model_summary": {
            "dataset_name": model_summary.get("dataset_name"),
            "random_rmse": model_summary.get("random_split", {}).get("rmse"),
            "scaffold_rmse": model_summary.get("scaffold_split", {}).get("rmse"),
            "temporal_rmse": model_summary.get("temporal_split", {}).get("rmse"),
        },
        "notebook_metrics": notebook_metrics,
        "top_candidates": top_candidate_records,
    }, previous_snapshot=previous_snapshot)
    context_notes.append(evolution_note)
    snapshot = {
        "run_label": run_label,
        "created_at": datetime.now().isoformat(timespec="seconds"),
        "candidate_source": candidate_source,
        "rl_top_candidates": (rl_candidate_df.head(5).to_dict(orient="records") if rl_candidate_df is not None else []),
        "model_summary": {
            "dataset_name": model_summary.get("dataset_name"),
            "random_rmse": model_summary.get("random_split", {}).get("rmse"),
            "scaffold_rmse": model_summary.get("scaffold_split", {}).get("rmse"),
            "temporal_rmse": model_summary.get("temporal_split", {}).get("rmse"),
        },
        "notebook_metrics": notebook_metrics,
        "top_candidates": top_candidate_records,
        "grid_image": str(grid_path) if grid_path else None,
        "top3_assets": top3_assets,
        "plots": copied_plots,
        "context_notes": context_notes,
        "evolution_note": evolution_note,
    }
    (run_dir / "snapshot.json").write_text(json.dumps(snapshot, indent=2), encoding="utf-8")
    return snapshot


def _load_history() -> list[dict]:
    if not HISTORY_INDEX.exists():
        return []
    return json.loads(HISTORY_INDEX.read_text(encoding="utf-8"))


def _save_history(history: list[dict]) -> None:
    HISTORY_DIR.mkdir(parents=True, exist_ok=True)
    HISTORY_INDEX.write_text(json.dumps(history, indent=2), encoding="utf-8")


def _append_context_memory(snapshot: dict) -> None:
    HISTORY_DIR.mkdir(parents=True, exist_ok=True)
    lines = [
        f"## {snapshot.get('run_label', 'unknown_run')}",
        f"- Created at: {snapshot.get('created_at', 'n/a')}",
        f"- Candidate source: {snapshot.get('candidate_source', 'n/a')}",
    ]
    for note in snapshot.get("context_notes", []):
        lines.append(f"- {note}")
    lines.append("")
    with CONTEXT_MEMORY.open("a", encoding="utf-8") as handle:
        handle.write("\n".join(lines) + "\n")


def _upsert_snapshot(history: list[dict], snapshot: dict) -> list[dict]:
    filtered = [entry for entry in history if entry.get("run_label") != snapshot.get("run_label")]
    filtered.append(snapshot)
    filtered.sort(key=lambda entry: entry.get("created_at", ""))
    return filtered


def _add_title(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = document.styles["Title"].font.size


def _add_bullet(document: Document, text: str) -> None:
    document.add_paragraph(text, style="List Bullet")


def _add_label(document: Document, label: str, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.add_run(f"{label}: ").bold = True
    paragraph.add_run(text)


def _figure_note(title: str, explanation: str, importance: str) -> dict[str, str]:
    return {
        "title": title,
        "explanation": explanation,
        "importance": importance,
    }


def _add_picture_if_exists(document: Document, path_str: str | None, width: float = 6.3) -> None:
    if not path_str:
        return
    path = Path(path_str)
    if not path.exists():
        return
    document.add_picture(str(path), width=Inches(width))


def _add_figure_with_notes(
    document: Document,
    path_str: str | None,
    figure_note: dict[str, str] | None,
    width: float = 6.3,
) -> None:
    if not path_str:
        return
    path = Path(path_str)
    if not path.exists():
        return
    title = figure_note.get("title") if figure_note else path.stem.replace("_", " ").title()
    document.add_paragraph(title)
    document.add_picture(str(path), width=Inches(width))
    if figure_note:
        _add_label(document, "What this figure shows", figure_note["explanation"])
        _add_label(document, "Why this figure matters", figure_note["importance"])


def _multi_agent_ablation_summary() -> dict[str, float] | None:
    ablation_df = _load_df(MULTI_AGENT_ABLATION)
    if ablation_df is None or ablation_df.empty:
        return None

    subset = ablation_df[ablation_df["strategy"].isin(["protected_final", "naive_proxy"])].copy()
    if subset.empty:
        return None

    common_topk = sorted(set(subset["top_k"].tolist()))
    if not common_topk:
        return None
    selected_topk = 100 if 100 in common_topk else common_topk[min(len(common_topk) - 1, 1)]

    protected = subset[(subset["strategy"] == "protected_final") & (subset["top_k"] == selected_topk)]
    naive = subset[(subset["strategy"] == "naive_proxy") & (subset["top_k"] == selected_topk)]
    if protected.empty or naive.empty:
        return None

    protected_row = protected.iloc[0]
    naive_row = naive.iloc[0]
    return {
        "top_k": float(selected_topk),
        "protected_pic50": float(protected_row.get("mean_predicted_pIC50", 0.0)),
        "naive_pic50": float(naive_row.get("mean_predicted_pIC50", 0.0)),
        "protected_risk": float(protected_row.get("mean_reward_hacking_risk", 0.0)),
        "naive_risk": float(naive_row.get("mean_reward_hacking_risk", 0.0)),
        "protected_pass": float(protected_row.get("audit_pass_rate", 0.0)),
        "naive_pass": float(naive_row.get("audit_pass_rate", 0.0)),
        "protected_review_fail": float(protected_row.get("review_or_fail_rate", 0.0)),
        "naive_review_fail": float(naive_row.get("review_or_fail_rate", 0.0)),
    }


def _build_figure_notes(model_summary: dict, notebook_metrics: dict) -> dict[str, dict[str, str]]:
    multiview_random = _fmt(model_summary.get("random_rmse"))
    multiview_scaffold = _fmt(model_summary.get("scaffold_rmse"))
    multiview_temporal = _fmt(model_summary.get("temporal_rmse"))
    audit_pass = _fmt(notebook_metrics.get("audit_pass_rate"))
    crossdb = _fmt(notebook_metrics.get("cross_database_mean_consensus"))
    feasibility = _fmt(notebook_metrics.get("mean_feasibility_score"))
    best_vina = _fmt(notebook_metrics.get("best_vina_affinity_kcal"))
    rediscovery = _fmt(notebook_metrics.get("rediscovery_protected_top10_recall"))
    source_holdout = _fmt(notebook_metrics.get("source_holdout_mean_rmse"))

    multi_agent_summary = _multi_agent_ablation_summary()
    if multi_agent_summary:
        multi_agent_explanation = (
            "This figure compares a reward-only single-agent selector against the protected multi-agent ranker across several shortlist sizes. "
            f"At top-{int(multi_agent_summary['top_k'])}, the protected system keeps mean predicted pIC50 at {_fmt(multi_agent_summary['protected_pic50'])} "
            f"versus {_fmt(multi_agent_summary['naive_pic50'])} for the naive selector, while reducing mean reward-hacking risk from "
            f"{_fmt(multi_agent_summary['naive_risk'])} to {_fmt(multi_agent_summary['protected_risk'])}."
        )
        multi_agent_importance = (
            "It matters because it makes the central project claim concrete: the multi-agent layer is not just cosmetic. "
            f"It preserves most of the potency signal while improving audit pass rate from {_fmt(multi_agent_summary['naive_pass'])} to {_fmt(multi_agent_summary['protected_pass'])} "
            f"and reducing review-or-fail cases from {_fmt(multi_agent_summary['naive_review_fail'])} to {_fmt(multi_agent_summary['protected_review_fail'])}."
        )
    else:
        multi_agent_explanation = (
            "This figure compares a reward-only single-agent selector against the protected multi-agent ranker across several shortlist sizes."
        )
        multi_agent_importance = (
            "It matters because it shows whether the audit layer improves shortlist credibility instead of simply reshuffling molecules."
        )

    return {
        "pipeline_flowchart.png": _figure_note(
            "Pipeline Flow Chart",
            "This flow chart summarizes the project from curated EGFR bioactivity data through multiview QSAR, candidate generation, protected ranking, structural and cross-database validation, and the final prospective batch.",
            "It matters because it helps judges see that the notebook is not one model plus one score. The project is a staged prioritization pipeline, and every stage removes a different failure mode.",
        ),
        "project_phase_capability_matrix.png": _figure_note(
            "Project Capability Matrix",
            "This matrix shows which major capabilities were present in each project phase, including dataset upgrades, audit logic, structural support, external validation, and prospective selection.",
            "It matters because it turns project evolution into something measurable: the system matured by adding safeguards and evidence layers, not only by adding more code.",
        ),
        "project_evolution_history.png": _figure_note(
            "Project Evolution Timeline",
            "This timeline tracks the main project checkpoints from the early Desktop baseline to the current evidence-aware platform, making it easy to connect each iteration to a concrete upgrade.",
            "It matters because competitions reward iteration quality. The figure shows that the current system is the result of successive design decisions rather than a one-shot build.",
        ),
        "literature_context_comparison.png": _figure_note(
            "Context Against Related EGFR Studies",
            "This comparison places OncoForge beside reported EGFR AI studies and industry-style baselines using public summary metrics and methodological context rather than claiming an exact apples-to-apples leaderboard.",
            "It matters because the notebook needs external anchors. The figure shows that the project's validation profile sits in a credible range while using broader audit and evidence safeguards than many simpler baselines.",
        ),
        "single_agent_vs_multi_agent.png": _figure_note(
            "Single-Agent vs Multi-Agent Ranking",
            multi_agent_explanation,
            multi_agent_importance,
        ),
        "risk_distribution_by_audit_status.png": _figure_note(
            "Audit Risk Distribution",
            "This histogram groups molecules by audit outcome and shows where reward-hacking risk accumulates across pass, review, and fail decisions.",
            "It matters because it demonstrates that the audit is selective rather than arbitrary: higher-risk molecules should concentrate in review or fail buckets if the protection logic is working.",
        ),
        "naive_vs_protected_scores.png": _figure_note(
            "Naive Score vs Protected Score",
            "This scatter plot compares the raw reward score with the final protected score, highlighting molecules most strongly demoted by the audit.",
            "It matters because it visualizes how the shortlist changes when the system stops trusting raw proxy optimization alone. A good protected pipeline should visibly separate suspicious high-score molecules from durable leads.",
        ),
        "audit_rank_demotions.png": _figure_note(
            "Largest Audit Demotions",
            "This bar chart shows which candidates lost the most positions after the anti-hacking audit reviewed them.",
            "It matters because it reveals where the protection system is actively preventing score inflation from dominating the final shortlist.",
        ),
        "top_leads_agent_support_heatmap.png": _figure_note(
            "Top-Lead Agent Support Heatmap",
            "This heatmap breaks the top protected candidates into support from potency, chemistry, safety, and domain agents.",
            "It matters because it shows whether the final leads are well balanced. Strong candidates should receive support from several perspectives, not from one narrow signal only.",
        ),
        "novelty_vs_applicability.png": _figure_note(
            "Novelty vs Applicability",
            "This plot shows how novelty interacts with applicability-domain support, with color indicating reward-hacking risk.",
            "It matters because novelty is useful only while the model still remains inside a believable evidence envelope. The best candidates balance originality with enough similarity to known chemistry.",
        ),
        "model_split_performance.png": _figure_note(
            "Validation Performance Across Splits",
            f"This figure compares model error and R2 under random, scaffold, and temporal validation. In the current snapshot, random RMSE is {multiview_random}, scaffold RMSE is {multiview_scaffold}, and temporal RMSE is {multiview_temporal}.",
            "It matters because scaffold and temporal splits are harder and more realistic than a simple random split. They tell the judges whether the model can generalize beyond close analogs.",
        ),
        "gpu_gnn_scaffold_benchmark.png": _figure_note(
            "GPU Graph Benchmark",
            "This figure compares GPU graph-based models with the classical multiview reference on scaffold validation.",
            "It matters because it shows that the project explored newer neural architectures, but still judged them by realistic validation rather than novelty alone.",
        ),
        "uncertainty_calibration.png": _figure_note(
            "Uncertainty Calibration",
            "This figure compares raw and calibrated uncertainty coverage to show whether the model's confidence estimates better match real prediction error after scaling.",
            "It matters because uncertainty is used downstream for selection and novelty control. Poorly calibrated confidence would make exploration look safer than it really is.",
        ),
        "structural_rescoring_scatter.png": _figure_note(
            "Structural Rescoring Support",
            "This scatter plot checks whether candidates that rank well in the protected pipeline also receive independent support from docking-based structural rescoring.",
            "It matters because the notebook should not rely on QSAR alone. Orthogonal structural support makes the final shortlist more convincing for follow-up work.",
        ),
        "vina_affinity_vs_priority.png": _figure_note(
            "Docking Strength vs Priority",
            "This figure compares AutoDock Vina affinity with the final structural or protected priority score for docked candidates.",
            "It matters because it shows whether stronger docking support is aligned with overall ranking rather than contradicting it. The best handoff candidates should not collapse when checked structurally.",
        ),
        "interaction_support_vs_vina.png": _figure_note(
            "Interaction Support vs Vina",
            "This plot looks beyond affinity alone by checking whether docked molecules also retain residue-level interaction support consistent with EGFR binding.",
            "It matters because a single docking score can be misleading. Interaction support adds a more interpretable structural layer.",
        ),
        "readiness_vs_structure.png": _figure_note(
            "Experimental Readiness vs Structural Support",
            "This figure compares structural support with the experimental-readiness score used to prepare a rational prospective batch.",
            "It matters because the project goal is not just ranking but handoff quality. Strong candidates should combine structure support with practical readiness for validation.",
        ),
        "cross_database_vs_potency.png": _figure_note(
            "Cross-Database Support vs Potency",
            "This scatter plot checks whether high predicted potency also appears in candidates that have stronger support across independent public databases.",
            "It matters because independent evidence is a key protection against overfitting to one source or one assay family.",
        ),
        "cross_database_consensus_vs_readiness.png": _figure_note(
            "Consensus vs Readiness",
            "This figure shows how cross-database consensus lines up with experimental readiness in the final candidate pool.",
            "It matters because the project aims to send forward molecules that are not only promising on paper but also supported enough to justify next-step testing.",
        ),
        "cross_database_status_counts.png": _figure_note(
            "Cross-Database Status Counts",
            "This chart counts how many candidates fall into strong, moderate, or weak independent-support categories.",
            "It matters because it converts a complex evidence story into a simple portfolio view, making it easier to judge how robust the shortlist really is.",
        ),
        "external_evidence_support_vs_potency.png": _figure_note(
            "External Evidence vs Potency",
            "This plot checks whether predicted potency remains supported when external evidence sources are considered separately from the main ranking model.",
            "It matters because the project becomes more credible when high-scoring molecules are also visible in orthogonal evidence channels.",
        ),
        "pubchem_assay_relevance.png": _figure_note(
            "PubChem Assay Relevance",
            "This bar chart ranks PubChem assays by their relevance to EGFR, showing which assays contribute the most meaningful external evidence.",
            "It matters because not every public assay is equally informative. The figure helps justify how PubChem was filtered into useful support instead of noisy background.",
        ),
        "feasibility_vs_potency.png": _figure_note(
            "Feasibility vs Potency",
            "This figure compares medicinal-chemistry feasibility with predicted potency for optimized candidates.",
            "It matters because a strong candidate must be more than potent. If potency rises only in synthetically unrealistic molecules, the project has not really improved the shortlist.",
        ),
        "generator_benchmark_overview.png": _figure_note(
            "Generator Benchmark Overview",
            "This multi-panel benchmark compares broad analog generation, AI-guided analogs, and iterative optimization on candidate count, priority, and audit quality.",
            "It matters because it shows which generation strategy actually improves the pipeline rather than just producing more molecules.",
        ),
        "rl_training_curve.png": _figure_note(
            "RL Training Curve",
            "This training curve shows how reward changes during the verifiable reinforcement-learning run.",
            "It matters because RL can look impressive even when it becomes unstable or exploits the reward. The curve helps show whether learning was controlled and interpretable.",
        ),
        "rl_reward_breakdown.png": _figure_note(
            "RL Reward Breakdown",
            "This figure decomposes reinforcement-learning reward into the components that drive the final agent behavior.",
            "It matters because it exposes the incentives given to the generator and makes reward design transparent instead of opaque.",
        ),
        "rl_external_evidence_vs_priority.png": _figure_note(
            "RL External Evidence vs Priority",
            "This plot checks whether RL-generated candidates that rank highly also carry stronger external evidence support.",
            "It matters because it tests whether RL is producing chemically supported ideas or merely learning to optimize a proxy score.",
        ),
        "gpu_rl_training_curve.png": _figure_note(
            "GPU RL Training Curve",
            "This curve tracks the GPU DQN training dynamics for the faster reinforcement-learning branch.",
            "It matters because it helps compare whether the faster branch improved efficiency without sacrificing stability or scientific plausibility.",
        ),
        "gpu_actor_critic_training_curve.png": _figure_note(
            "GPU Actor-Critic Training Curve",
            "This figure shows the optimization trajectory of the actor-critic branch explored as an advanced extension.",
            "It matters because it documents that multiple RL formulations were tested and assessed rather than reported selectively.",
        ),
        "model_robustness_scaffold.png": _figure_note(
            "Scaffold Robustness Benchmark",
            "This benchmark summarizes scaffold performance across model families and repeated runs to show variability, not just the best single result.",
            "It matters because robust science depends on stability. A model that wins once but varies widely is weaker for real lead prioritization.",
        ),
        "source_holdout_rmse.png": _figure_note(
            "Source Holdout RMSE",
            f"This figure reports error when one public source is left out of training and used only for testing. The current mean holdout RMSE is {source_holdout}.",
            "It matters because it tests transfer across database boundaries, which is much closer to real-world deployment than validating only on mixed pooled data.",
        ),
        "source_holdout_recall.png": _figure_note(
            "Source Holdout Recall",
            "This figure measures how well the ranking still recovers strong candidates when an entire source is withheld from training.",
            "It matters because recall under source shift is a practical test of whether the model preserves useful prioritization power outside its original evidence mix.",
        ),
        "rediscovery_recall_at_k.png": _figure_note(
            "Rediscovery Recall at K",
            f"This curve measures how often the protected ranker recovers known strong molecules inside the top-k window. The current protected top-10 recall is {rediscovery}.",
            "It matters because rediscovery is an intuitive sanity check: a strong ranking system should still find high-value known chemistry in a difficult benchmark panel.",
        ),
        "rediscovery_rank_shift.png": _figure_note(
            "Rediscovery Rank Shift",
            "This figure shows how the positions of known actives move when the system switches from naive ranking to protected ranking.",
            "It matters because it reveals whether the audit hurts useful molecules or mainly removes suspicious proxy exploiters.",
        ),
        "challenge_rank_shift.png": _figure_note(
            "Reward-Hacking Challenge Rank Shift",
            "This figure tracks how deliberately constructed proxy-exploit molecules are pushed down after audit in the challenge benchmark.",
            "It matters because it is direct evidence that the anti-hacking logic responds to adversarial chemistry instead of only to normal candidates.",
        ),
        "challenge_status_rates.png": _figure_note(
            "Reward-Hacking Challenge Status Rates",
            "This chart compares audit outcomes for trusted controls and proxy-exploit cohorts in the challenge benchmark.",
            "It matters because it shows whether the audit is discriminative: trusted controls should pass more often than intentionally suspicious molecules.",
        ),
        "prospective_batch_readiness_vs_novelty.png": _figure_note(
            "Prospective Batch: Readiness vs Novelty",
            "This figure maps the selected prospective batch along readiness and novelty so that the final handoff is not dominated by only safe analogs or only speculative ideas.",
            "It matters because the output of the project is a balanced portfolio, not just the highest scalar score.",
        ),
        "marketed_vs_generated_boxplots.png": _figure_note(
            "Marketed vs Generated Benchmark",
            "This boxplot compares generated candidates with marketed or benchmark EGFR molecules on potency, QED, and reward-hacking risk.",
            "It matters because it grounds the shortlist against molecules that already define the target space, making the performance story easier to interpret for judges.",
        ),
        "structural_benchmark_boxplots.png": _figure_note(
            "Structural Benchmark Boxplots",
            "This figure compares marketed, optimized, generated, and prospective molecules on structural metrics such as Vina affinity and interaction support.",
            "It matters because it shows where the generated chemistry stands relative to known EGFR drugs and internal optimized leads after structural checking.",
        ),
        "technical_notebook_chemical_space.png": _figure_note(
            "Chemical Space Snapshot",
            "This PCA projection places ranked leads, marketed EGFR molecules, and the novel shortlist in the same reduced chemical space.",
            "It matters because it visually explains the balance between familiarity and novelty. The project should explore beyond marketed chemistry without drifting into unsupported space.",
        ),
    }


def _add_run_comparison(document: Document, history: list[dict]) -> None:
    if not history:
        return
    document.add_heading("Run-to-Run Comparison", level=1)
    table = document.add_table(rows=1, cols=6)
    header = table.rows[0].cells
    header[0].text = "Run"
    header[1].text = "Scaffold RMSE"
    header[2].text = "Mean Feasibility"
    header[3].text = "Best Vina"
    header[4].text = "Top Score"
    header[5].text = "Top pIC50"

    for snapshot in history:
        row = table.add_row().cells
        metrics = snapshot.get("notebook_metrics", {})
        top_candidate = (snapshot.get("top_candidates") or [{}])[0]
        row[0].text = str(snapshot.get("run_label", "n/a"))
        row[1].text = _fmt(snapshot.get("model_summary", {}).get("scaffold_rmse"))
        row[2].text = _fmt(metrics.get("mean_feasibility_score"))
        row[3].text = _fmt(metrics.get("best_vina_affinity_kcal"))
        row[4].text = _fmt(top_candidate.get("final_score"))
        row[5].text = _fmt(top_candidate.get("predicted_pIC50"))


def _competition_context() -> dict:
    context = _load_json(COMPETITION_CONTEXT_JSON)
    return context if isinstance(context, dict) else {}


def _add_isef_front_matter(document: Document, history: list[dict], figure_notes: dict[str, dict[str, str]]) -> None:
    latest = history[-1] if history else {}
    latest_metrics = latest.get("notebook_metrics", {})
    latest_model = latest.get("model_summary", {})
    latest_top = (latest.get("top_candidates") or [{}])[0]
    competition_context = _competition_context()

    dataset_size = int(competition_context.get("model_dataset_size", 0) or 0)
    ranked_molecules = int(competition_context.get("ranked_molecules", 0) or 0)
    prospective_batch_size = int(competition_context.get("prospective_batch_size", 0) or 0)

    document.add_heading("Abstract", level=1)
    document.add_paragraph(
        "OncoForge is an ISEF-style computational drug-discovery project centered on EGFR lead optimization. "
        "The system is designed to prioritize chemically plausible, evidence-supported candidate molecules rather than to claim finished drug discovery. "
        "Its main novelty is the combination of a classical multiview QSAR model, a protected multi-agent ranking layer, orthogonal structural rescoring, "
        "cross-database evidence checks, feasibility and experimental-readiness scoring, and reinforcement-learning branches whose outputs are audited against reward hacking."
    )
    document.add_paragraph(
        f"In the current snapshot, the model uses {dataset_size} cleaned training molecules and ranks {ranked_molecules} protected candidates. "
        f"The current validation profile is random RMSE {_fmt(latest_model.get('random_rmse'))}, scaffold RMSE {_fmt(latest_model.get('scaffold_rmse'))}, "
        f"and temporal RMSE {_fmt(latest_model.get('temporal_rmse'))}. The optimized-candidate ecosystem reaches mean feasibility {_fmt(latest_metrics.get('mean_feasibility_score'))}, "
        f"best Vina affinity {_fmt(latest_metrics.get('best_vina_affinity_kcal'))} kcal/mol, and a prospective validation batch of {prospective_batch_size} candidates. "
        f"The leading protected candidate in this run has final score {_fmt(latest_top.get('final_score'))} and predicted pIC50 {_fmt(latest_top.get('predicted_pIC50'))}."
    )

    document.add_heading("Research Problem", level=1)
    _add_bullet(document, "The practical challenge is that EGFR chemical space is vast, while wet-lab validation is expensive and limited.")
    _add_bullet(document, "A naive optimization pipeline can over-trust proxy scores, extrapolate outside evidence-supported chemistry, or reward-hack itself into unrealistic molecules.")
    _add_bullet(document, "The project goal is therefore to rank molecules for follow-up using stronger safeguards, not just larger candidate counts.")

    document.add_heading("Hypothesis and Design Logic", level=1)
    _add_bullet(document, "If potency prediction is separated from chemistry, safety, applicability, and audit logic, the final shortlist becomes more credible than a single-score ranker.")
    _add_bullet(document, "If independent public databases and structural rescoring are layered on top of QSAR predictions, the project can better distinguish plausible leads from score-exploiting artifacts.")
    _add_bullet(document, "If the shortlist is selected as a prospective batch instead of a top-N list, the final portfolio can balance exploitation, novelty, uncertainty, and diversity.")

    document.add_heading("Method Overview", level=1)
    _add_bullet(document, "Data: ChEMBL-centered EGFR training set upgraded to a multisource evidence ecosystem using BindingDB, Papyrus, ExCAPE-DB, PubChem, and Guide to Pharmacology.")
    _add_bullet(document, "Prediction: classical multiview ensemble with random, scaffold, and temporal validation rather than a single optimistic split.")
    _add_bullet(document, "Protection: a multi-agent ranker combines potency, chemistry quality, safety, novelty, applicability-domain, and anti-reward-hacking logic.")
    _add_bullet(document, "Orthogonal support: docking, residue-level interaction analysis, feasibility evidence, and experimental-readiness scoring.")
    _add_bullet(document, "Generation: broad analog expansion, AI-guided analogs, iterative optimization, and RL branches with traceable medicinal-chemistry moves.")
    _add_bullet(document, "Decision output: market comparison, rediscovery checks, source holdout, reward-hacking challenge, and a prospective validation batch.")
    document.add_heading("Pipeline Overview Figure", level=2)
    _add_figure_with_notes(
        document,
        str(NOTEBOOK_DIR / "pipeline_flowchart.png"),
        figure_notes.get("pipeline_flowchart.png"),
        width=6.8,
    )

    document.add_heading("Project Evolution", level=1)
    document.add_paragraph(
        "The project did not appear in its final form. It evolved from a baseline Desktop version focused on single-source QSAR and analog generation into an evidence-aware lead-prioritization platform."
    )
    for phase in PROJECT_PHASES:
        document.add_heading(f"{phase.phase_id} | {phase.title}", level=2)
        _add_label(document, "Date / commit", f"{phase.date_label} / {phase.commit}")
        _add_label(document, "Main focus", phase.focus)
        for upgrade in phase.upgrades:
            _add_bullet(document, upgrade)

    document.add_heading("Evolution Figures", level=2)
    _add_figure_with_notes(
        document,
        str(NOTEBOOK_DIR / "project_phase_capability_matrix.png"),
        figure_notes.get("project_phase_capability_matrix.png"),
        width=6.8,
    )
    _add_figure_with_notes(
        document,
        str(NOTEBOOK_DIR / "project_evolution_history.png"),
        figure_notes.get("project_evolution_history.png"),
        width=6.8,
    )

    document.add_heading("Current Competition Strengths", level=1)
    _add_bullet(document, "Protected ranking matters: the anti-hacking audit demotes suspicious molecules and changes the final shortlist substantially.")
    _add_bullet(document, "Generalization is examined with scaffold split, temporal split, and leave-one-source-out benchmarking rather than one convenient split.")
    _add_bullet(document, "Rediscovery and source-holdout benchmarks create a stronger scientific credibility story than raw generation counts alone.")
    _add_bullet(document, "The prospective validation batch is closer to a real experimental handoff than a generic top-scoring table.")
    _add_bullet(document, "The project includes a large technical glossary and automated report generation, which improves explainability and reproducibility for competition review.")
    document.add_heading("Core Ranking Figure", level=2)
    _add_figure_with_notes(
        document,
        str(NOTEBOOK_DIR / "single_agent_vs_multi_agent.png"),
        figure_notes.get("single_agent_vs_multi_agent.png"),
        width=6.8,
    )

    document.add_heading("Context Against Related Studies and Industry", level=1)
    document.add_paragraph(
        "The following comparison plots are contextual rather than strict leaderboards. The cited studies differ in endpoint type, split strategy, chemical scope, and objective, "
        "so the safe claim is that OncoForge sits in a credible performance range for related EGFR AI studies while using broader evidence-aware validation than many narrower baselines."
    )
    _add_figure_with_notes(
        document,
        str(NOTEBOOK_DIR / "literature_context_comparison.png"),
        figure_notes.get("literature_context_comparison.png"),
        width=6.8,
    )
    _add_figure_with_notes(
        document,
        str(NOTEBOOK_DIR / "marketed_vs_generated_boxplots.png"),
        figure_notes.get("marketed_vs_generated_boxplots.png"),
        width=6.8,
    )
    _add_figure_with_notes(
        document,
        str(NOTEBOOK_DIR / "structural_benchmark_boxplots.png"),
        figure_notes.get("structural_benchmark_boxplots.png"),
        width=6.8,
    )

    document.add_heading("Limitations and Honest Boundaries", level=1)
    _add_bullet(document, "This is an in-silico prioritization system, not proof of clinical efficacy or a claim of completed drug discovery.")
    _add_bullet(document, "The temporal split remains difficult, which is expected for chemistry published later than the training evidence.")
    _add_bullet(document, "The GPU neural branches are useful as research extensions, but the strongest empirical story in this codebase still comes from protected classical ranking plus external evidence layers.")
    _add_bullet(document, "The software test suite is smoke-level, so the scientific rigor is stronger than the software-coverage story.")

    document.add_heading("Selected Literature and Quotations", level=1)
    document.add_paragraph(
        "The notebook uses a deliberately mixed reference set: foundational cheminformatics papers, public database papers, molecular-generation references, and EGFR-specific comparison studies."
    )
    for entry in COMPETITION_LITERATURE:
        document.add_heading(entry.title, level=2)
        _add_label(document, "Citation", entry.citation)
        _add_label(document, "Why it matters here", entry.why_it_matters)
        _add_label(document, "Short quote", f"\"{entry.short_quote}\"")
        _add_label(document, "Source", entry.url)
        if entry.comparison_note:
            _add_label(document, "Comparison note", entry.comparison_note)

    document.add_heading("Glossary Companion", level=1)
    document.add_paragraph(
        "A dedicated companion document, OncoForge_Buzzword_Glossary.docx, explains the project vocabulary in plain language. "
        "That glossary should be treated as the technical dictionary for judges or team members who need rapid onboarding."
    )


def _build_docx(history: list[dict]) -> Path:
    document = Document()
    latest_snapshot = history[-1] if history else {}
    figure_notes = _build_figure_notes(
        latest_snapshot.get("model_summary", {}),
        latest_snapshot.get("notebook_metrics", {}),
    )
    _add_title(
        document,
        "OncoForge: ISEF-Style Technical Notebook for an Audited EGFR Lead-Optimization Platform",
    )
    document.add_paragraph(
        "Technical notebook generated automatically from major pipeline runs. "
        "This document combines ISEF-style front matter, project evolution, literature positioning, and detailed run-by-run appendices."
    )
    _add_isef_front_matter(document, history, figure_notes)
    _add_run_comparison(document, history)

    latest_run_label = latest_snapshot.get("run_label")
    for snapshot in history:
        document.add_section(WD_SECTION.NEW_PAGE)
        document.add_heading(f"Run {snapshot['run_label']}", level=1)
        document.add_paragraph(f"Created at: {snapshot.get('created_at', 'n/a')}")
        document.add_paragraph(f"Candidate source: {snapshot.get('candidate_source', 'n/a')}")

        model_summary = snapshot.get("model_summary", {})
        document.add_heading("Run Summary", level=2)
        _add_bullet(document, f"Random RMSE: {model_summary.get('random_rmse', 'n/a')}")
        _add_bullet(document, f"Scaffold RMSE: {model_summary.get('scaffold_rmse', 'n/a')}")
        _add_bullet(document, f"Temporal RMSE: {model_summary.get('temporal_rmse', 'n/a')}")
        notebook_metrics = snapshot.get("notebook_metrics", {})
        _add_bullet(document, f"Audit pass rate: {notebook_metrics.get('audit_pass_rate', 'n/a')}")
        _add_bullet(document, f"Audit fail rate: {notebook_metrics.get('audit_fail_rate', 'n/a')}")
        _add_bullet(document, f"Median reward hacking risk: {notebook_metrics.get('median_reward_hacking_risk', 'n/a')}")
        _add_bullet(document, f"Mean feasibility score: {notebook_metrics.get('mean_feasibility_score', 'n/a')}")
        _add_bullet(document, f"Mean experimental readiness: {notebook_metrics.get('mean_experimental_readiness', 'n/a')}")
        _add_bullet(document, f"Evidence arbiter mean support: {notebook_metrics.get('evidence_arbiter_mean_support', 'n/a')}")
        _add_bullet(document, f"Evidence arbiter pass rate: {notebook_metrics.get('evidence_arbiter_pass_rate', 'n/a')}")
        _add_bullet(document, f"Cross-database mean consensus: {notebook_metrics.get('cross_database_mean_consensus', 'n/a')}")
        _add_bullet(document, f"Cross-database strong rate: {notebook_metrics.get('cross_database_strong_rate', 'n/a')}")
        _add_bullet(document, f"External evidence mean support: {notebook_metrics.get('external_evidence_mean_support', 'n/a')}")
        _add_bullet(document, f"Papyrus molecules: {notebook_metrics.get('papyrus_unique_molecules', 'n/a')}")
        _add_bullet(document, f"ExCAPE molecules: {notebook_metrics.get('excape_unique_molecules', 'n/a')}")
        _add_bullet(document, f"PubChem mean enriched evidence: {notebook_metrics.get('pubchem_mean_enriched_evidence_score', 'n/a')}")
        _add_bullet(document, f"PubChem strong evidence rate: {notebook_metrics.get('pubchem_strong_evidence_rate', 'n/a')}")
        _add_bullet(document, f"RL mean external evidence support: {notebook_metrics.get('rl_mean_external_evidence_support', 'n/a')}")
        _add_bullet(document, f"RL ready rate: {notebook_metrics.get('rl_readiness_ready_rate', 'n/a')}")
        _add_bullet(document, f"GPU GNN best scaffold model: {notebook_metrics.get('gpu_gnn_best_scaffold_model', 'n/a')}")
        _add_bullet(document, f"GPU GNN best scaffold RMSE: {notebook_metrics.get('gpu_gnn_best_scaffold_rmse', 'n/a')}")
        _add_bullet(document, f"GPU RL mean external evidence support: {notebook_metrics.get('gpu_rl_mean_external_evidence_support', 'n/a')}")
        _add_bullet(document, f"GPU RL best episode return: {notebook_metrics.get('gpu_rl_best_episode_return', 'n/a')}")
        _add_bullet(document, f"Source holdout mean RMSE: {notebook_metrics.get('source_holdout_mean_rmse', 'n/a')}")
        _add_bullet(document, f"Source holdout best source: {notebook_metrics.get('source_holdout_best_source', 'n/a')}")
        _add_bullet(document, f"Rediscovery protected top-10 recall: {notebook_metrics.get('rediscovery_protected_top10_recall', 'n/a')}")
        _add_bullet(document, f"Rediscovery protected top-20 recall: {notebook_metrics.get('rediscovery_protected_top20_recall', 'n/a')}")
        _add_bullet(document, f"Best Vina affinity: {notebook_metrics.get('best_vina_affinity_kcal', 'n/a')}")
        _add_bullet(document, f"Prospective batch size: {notebook_metrics.get('prospective_batch_size', 'n/a')}")

        document.add_heading("Context Notes", level=2)
        if snapshot.get("evolution_note"):
            document.add_paragraph(snapshot["evolution_note"])
        for note in snapshot.get("context_notes", []):
            if note == snapshot.get("evolution_note"):
                continue
            _add_bullet(document, note)

        document.add_heading("Top Molecules", level=2)
        _add_picture_if_exists(document, snapshot.get("grid_image"))
        for idx, candidate in enumerate(snapshot.get("top_candidates", []), start=1):
            document.add_paragraph(
                f"Top {idx}: score={_fmt(candidate.get('final_score'))}, "
                f"pIC50={_fmt(candidate.get('predicted_pIC50'))}, "
                f"QED={_fmt(candidate.get('QED'))}"
            )
            if "feasibility_score" in candidate:
                document.add_paragraph(f"Feasibility score: {_fmt(candidate.get('feasibility_score'))}")
            if "vina_affinity_kcal" in candidate:
                document.add_paragraph(f"Vina affinity: {_fmt(candidate.get('vina_affinity_kcal'))} kcal/mol")
            document.add_paragraph(candidate.get("smiles", ""))

        if snapshot.get("rl_top_candidates"):
            document.add_heading("Verifiable RL Snapshot", level=2)
            for idx, candidate in enumerate(snapshot.get("rl_top_candidates", [])[:3], start=1):
                document.add_paragraph(
                    f"RL {idx}: score={_fmt(candidate.get('rl_priority_score'))}, "
                    f"pIC50={_fmt(candidate.get('predicted_pIC50'))}, "
                    f"cross-db={_fmt(candidate.get('cross_database_consensus_score'))}, "
                    f"external evidence={_fmt(candidate.get('external_evidence_support'))}, "
                    f"readiness={_fmt(candidate.get('experimental_readiness_score'))}"
                )
                document.add_paragraph(candidate.get("smiles", ""))

        document.add_heading("3D Candidate Views", level=2)
        for asset in snapshot.get("top3_assets", []):
            summary = asset.get("summary", {})
            document.add_paragraph(
                f"Candidate rank {asset.get('rank', '?')} | "
                f"score={_fmt(summary.get('final_score'))} | "
                f"pIC50={_fmt(summary.get('predicted_pIC50'))} | "
                f"QED={_fmt(summary.get('QED'))}"
            )
            if "docking_rescore" in summary:
                document.add_paragraph(f"Structural rescoring support: {_fmt(summary.get('docking_rescore'))}")
            if "vina_affinity_kcal" in summary:
                document.add_paragraph(f"Vina affinity: {_fmt(summary.get('vina_affinity_kcal'))} kcal/mol")
            if "feasibility_score" in summary:
                document.add_paragraph(f"Feasibility score: {_fmt(summary.get('feasibility_score'))}")
            document.add_paragraph(summary.get("smiles", ""))
            _add_picture_if_exists(document, asset.get("image"), width=6.6)

        document.add_heading("Run Visualizations", level=2)
        if snapshot.get("run_label") == latest_run_label:
            document.add_paragraph(
                "The current run includes figure-by-figure interpretation so the notebook reads more like a classical technical report and less like a raw image appendix."
            )
        for plot_path in snapshot.get("plots", []):
            plot_filename = Path(plot_path).name
            note = figure_notes.get(plot_filename)
            if snapshot.get("run_label") == latest_run_label:
                _add_figure_with_notes(document, plot_path, note)
            else:
                plot_name = plot_filename.replace("_", " ").replace(".png", "")
                document.add_paragraph(plot_name.title())
                _add_picture_if_exists(document, plot_path)

    DOCX_PATH.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(DOCX_PATH))
    ENGLISH_ALIAS_PATH.write_bytes(DOCX_PATH.read_bytes())
    return DOCX_PATH


def main(argv: list[str] | None = None) -> None:
    run_label = _current_run_label()
    snapshot = _prepare_run_snapshot(run_label)
    history = _load_history()
    history = _upsert_snapshot(history, snapshot)
    _save_history(history)
    _append_context_memory(snapshot)
    out_path = _build_docx(history)
    print(f"[OK] Saved technical notebook Word document: {out_path}")
    print(f"[OK] Updated run history: {HISTORY_INDEX}")


if __name__ == "__main__":
    main()
